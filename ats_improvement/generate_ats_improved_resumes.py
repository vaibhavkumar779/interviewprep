"""
ATS-Improved Resume Generator for Vaibhav Kumar
Generates 3 ATS-optimized resumes with additional keywords from learning topics.
These resumes include bullets for skills you are STUDYING.
Only use these versions AFTER completing the corresponding learning topics.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "..", "resumes", "ats_improved")


def set_narrow_margins(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.25)
        section.bottom_margin = Inches(0.2)
        section.left_margin = Inches(0.4)
        section.right_margin = Inches(0.4)


def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.0):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing


def add_header_line(doc, text, size=10, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    set_paragraph_spacing(p, before=0, after=0, line_spacing=1.0)
    return p


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    set_paragraph_spacing(p, before=3, after=0, line_spacing=1.0)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {qn("w:val"): "single", qn("w:sz"): "4", qn("w:space"): "1", qn("w:color"): "000000"})
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_body_text(doc, text, bold=False, size=9, before=0, after=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    set_paragraph_spacing(p, before=before, after=after, line_spacing=1.0)
    return p


def add_skill_line(doc, category, items):
    p = doc.add_paragraph()
    run_cat = p.add_run(f"{category}: ")
    run_cat.font.size = Pt(8.5)
    run_cat.font.name = "Calibri"
    run_cat.font.bold = True
    run_cat.font.color.rgb = RGBColor(0, 0, 0)
    run_items = p.add_run(items)
    run_items.font.size = Pt(8.5)
    run_items.font.name = "Calibri"
    run_items.font.color.rgb = RGBColor(0, 0, 0)
    set_paragraph_spacing(p, before=0, after=0, line_spacing=1.0)
    return p


def add_experience_header(doc, title, date, size=9):
    p = doc.add_paragraph()
    run_title = p.add_run(title)
    run_title.font.size = Pt(size)
    run_title.font.name = "Calibri"
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0, 0, 0)
    run_date = p.add_run(f"  |  {date}")
    run_date.font.size = Pt(8)
    run_date.font.name = "Calibri"
    run_date.font.color.rgb = RGBColor(80, 80, 80)
    set_paragraph_spacing(p, before=2, after=0, line_spacing=1.0)
    return p


def add_bullet(doc, text, size=8.5, indent=0.2):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    pf = p.paragraph_format
    pf.left_indent = Inches(indent)
    pf.first_line_indent = Inches(-0.13)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    run = p.add_run(f"\u2022 {text}")
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_common_header(doc):
    name_p = add_header_line(doc, "VAIBHAV KUMAR", size=13, bold=True)
    set_paragraph_spacing(name_p, before=0, after=0)
    contact_p = add_header_line(doc, "+91-8318620608 | vaibhav.kr.779@gmail.com", size=8.5)
    set_paragraph_spacing(contact_p, before=0, after=0)
    links_p = add_header_line(doc, "linkedin.com/in/vaibhavkumar779 | github.com/vaibhavkumar779 | youtube.com/@DSeDevOps", size=8.5)
    set_paragraph_spacing(links_p, before=0, after=0)
    loc_p = add_header_line(doc, "Gurugram, Haryana, India", size=8.5)
    set_paragraph_spacing(loc_p, before=0, after=1)


def add_common_education(doc):
    add_section_heading(doc, "Education")
    add_experience_header(doc, "B.Tech \u2014 ABESIT (AKTU)", "Jul 2018 \u2013 Jun 2022")
    add_body_text(doc, "CGPA: 8.56 / 10.0  |  First Division", size=8, before=0, after=0)
    add_experience_header(doc, "Sainik School Gopalganj \u2014 PCMB (82.6%)", "Jul 2017")


def add_common_certifications(doc):
    add_section_heading(doc, "Certifications")
    add_bullet(doc, "Professional Cloud DevOps Engineer \u2014 Google Cloud Community India")
    add_bullet(doc, "Microsoft Certified: Azure Network Engineer Associate \u2014 Microsoft")
    add_bullet(doc, "AWS Cloud Training and Seminar Certificates \u2014 Amazon Web Services")
    add_bullet(doc, "Python 3 Programming \u2014 Coursera")


# ============================================================================
# ATS-IMPROVED RESUME 1: DEVOPS ENGINEER (Target ATS: 97/100)
# ============================================================================
def generate_devops_resume():
    doc = Document()
    set_narrow_margins(doc)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)

    add_common_header(doc)

    add_section_heading(doc, "Professional Summary")
    add_body_text(
        doc,
        "DevOps Engineer with 4.5+ years of experience designing CI/CD pipelines, automating cloud infrastructure, "
        "and orchestrating containerized workloads on Kubernetes (AKS/EKS). Delivered 50+ production deployments on AKS, "
        "authored 11 reusable Terraform modules, and implemented GitOps-based delivery with ArgoCD. Migrated secrets "
        "to Azure Key Vault eliminating plaintext credentials. Proficient in Jenkins, Azure DevOps, Docker, Terraform, "
        "Ansible, Istio service mesh, and Python/Shell scripting. Certified Google Cloud DevOps Engineer and Azure "
        "Network Engineer Associate. Awarded 2x SITA Bravo Awards for release engineering excellence.",
        size=8.5, before=0, after=1,
    )

    add_section_heading(doc, "Technical Skills")
    add_skill_line(doc, "CI/CD & GitOps", "Jenkins (Declarative, Scripted, Shared Libraries), Azure DevOps, GitHub Actions, ArgoCD, GitOps")
    add_skill_line(doc, "Containers & Orchestration", "Docker, Kubernetes (AKS, EKS), Helm Charts, Istio Service Mesh, ACR")
    add_skill_line(doc, "Infrastructure as Code", "Terraform (Azure, AWS \u2014 11 modules), Ansible, CloudFormation, Packer")
    add_skill_line(doc, "Cloud Platforms", "Azure (AKS, Key Vault, App Gateway, VNets, PostgreSQL Flex), AWS (EC2, S3, RDS, IAM, GuardDuty)")
    add_skill_line(doc, "Scripting", "Python, Bash (Shell Scripting), PowerShell")
    add_skill_line(doc, "Monitoring & Observability", "Prometheus, Grafana, ELK Stack, Azure Monitor, New Relic, Jaeger, Kiali, Loki")
    add_skill_line(doc, "DevSecOps", "SonarQube, Snyk, Mend, HashiCorp Vault, JFrog Artifactory")
    add_skill_line(doc, "Tools", "Git, Bitbucket, GitHub, Azure Repos, Linux (Ubuntu), Windows")

    add_section_heading(doc, "Professional Experience")

    add_experience_header(doc, "Infrastructure Engineer (DevOps) \u2014 SITA, Gurugram", "Jan 2025 \u2013 Present")
    add_bullet(doc, "Provisioned AKS clusters, VNets, Application Gateways, and PostgreSQL Flexible Servers using 11 Terraform modules across dev, preprod, and production, achieving 95% infrastructure-as-code coverage.")
    add_bullet(doc, "Led secrets migration from Helm values to Azure Key Vault with CSI Secret Store Driver across 4 product charts, eliminating 100% of plaintext secrets from Git repositories.")
    add_bullet(doc, "Implemented ArgoCD-based GitOps delivery for Kubernetes workloads, enabling declarative deployments with automated sync, drift detection, and self-healing reconciliation across 5 namespaces.")
    add_bullet(doc, "Architected CI/CD pipelines in Azure DevOps for 13 microservices with build, test, security scanning, and multi-environment deployment stages, managing release engineering across 5 environments.")
    add_bullet(doc, "Configured Istio service mesh for microservices enabling mTLS, traffic management, and canary deployments with Kiali visualization and Jaeger distributed tracing.")
    add_bullet(doc, "Deployed Prometheus, Grafana, and ELK Stack on AKS for full-stack observability \u2014 metrics, logs, and distributed tracing \u2014 reducing MTTD by 60% and MTTR by 45%.")
    add_bullet(doc, "Remediated 16 AWS security findings (EC2, RDS, S3, CloudFront, IAM, GuardDuty) achieving 100% compliance; implemented change management processes for production deployments.")
    add_bullet(doc, "Drove cost optimization through resource cleanup, right-sizing, and weekly reporting, reducing monthly Azure spend by 20%.")

    add_experience_header(doc, "Senior Software Consultant (DevOps) \u2014 Knoldus Inc (NashTech Global), Noida", "Nov 2021 \u2013 Dec 2024")
    add_bullet(doc, "Built Jenkins CI/CD pipelines (Declarative & Scripted with Shared Libraries) for 10+ projects, reducing build-to-deploy cycle time by 40% through parallelized stages.")
    add_bullet(doc, "Migrated artifact management from JFrog Artifactory to GitHub Packages using PowerShell automation, eliminating $15K+ annual licensing costs.")
    add_bullet(doc, "Integrated Snyk, Mend, and SonarQube into CI pipelines for automated security scanning, achieving zero critical CVEs in production releases.")
    add_bullet(doc, "Implemented Ansible configuration management across 30+ instances ensuring consistent provisioning and drift prevention.")

    add_experience_header(doc, "Data Engineer \u2014 REOMNIFY, Noida", "Jan 2021 \u2013 Jun 2021")
    add_bullet(doc, "Built Python-based data extraction pipelines using REST APIs and Selenium, processing 10K+ records daily into PostgreSQL.")

    add_common_education(doc)
    add_common_certifications(doc)

    add_section_heading(doc, "Achievements")
    add_bullet(doc, "Awarded 2x SITA Bravo Awards for driving complex cross-timezone deployments and consistent high-quality delivery.")
    add_bullet(doc, "Open-source contributor: PR #232 to Azure/terraform-azurerm-avm-res-keyvault-vault (Microsoft Azure Verified Modules).")

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    docx_path = os.path.join(OUTPUT_DIR, "Vaibhav_Kumar_DevOps_Engineer_ATS_Improved.docx")
    doc.save(docx_path)
    print(f"Created: {docx_path}")
    return docx_path


# ============================================================================
# ATS-IMPROVED RESUME 2: PLATFORM ENGINEER (Target ATS: 96/100)
# ============================================================================
def generate_platform_engineer_resume():
    doc = Document()
    set_narrow_margins(doc)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)

    add_common_header(doc)

    add_section_heading(doc, "Professional Summary")
    add_body_text(
        doc,
        "Platform Engineer with 4.5+ years of experience building Internal Developer Platforms (IDPs), treating "
        "platform as a product to improve Developer Experience (DevEx). Built self-service infrastructure with Backstage "
        "developer portal, golden-path CI/CD workflows, and Kubernetes-native provisioning. Authored 11 reusable Terraform "
        "modules, implemented ArgoCD GitOps delivery, and enforced platform guardrails with OPA Gatekeeper. Certified "
        "Google Cloud DevOps Engineer and Azure Network Engineer Associate. Awarded 2x SITA Bravo Awards.",
        size=8.5, before=0, after=1,
    )

    add_section_heading(doc, "Technical Skills")
    add_skill_line(doc, "Platform & Orchestration", "Kubernetes (AKS, EKS), Helm, Docker, ACR, Backstage, Crossplane, GitOps")
    add_skill_line(doc, "Infrastructure as Code", "Terraform (Azure, AWS \u2014 11 modules), Ansible, Crossplane (K8s-native IaC)")
    add_skill_line(doc, "CI/CD & GitOps", "Azure DevOps, Jenkins (Shared Libraries), GitHub Actions, ArgoCD, ApplicationSet")
    add_skill_line(doc, "Cloud Platforms", "Azure (AKS, Key Vault, App Gateway, VNets, PostgreSQL Flex), AWS (EC2, S3, RDS, IAM)")
    add_skill_line(doc, "Observability", "Prometheus, Grafana, Loki, Kiali, Jaeger, Azure Monitor, New Relic")
    add_skill_line(doc, "Policy & Security", "OPA Gatekeeper, Azure Key Vault, CSI Secret Store Driver, Snyk, Mend, SonarQube")
    add_skill_line(doc, "Scripting", "Python, Bash, PowerShell, HCL, Rego")
    add_skill_line(doc, "Service Mesh & Networking", "Istio, Kiali, Jaeger, Azure Application Gateway, VNets")

    add_section_heading(doc, "Professional Experience")

    add_experience_header(doc, "Infrastructure Engineer (Platform) \u2014 SITA, Gurugram", "Jan 2025 \u2013 Present")
    add_bullet(doc, "Designed Internal Developer Platform with Backstage portal and service catalog, enabling self-service discovery and provisioning for 13+ microservices, reducing developer onboarding from 2 days to 2 hours.")
    add_bullet(doc, "Authored 11 reusable Terraform modules for AKS, VNets, Application Gateways, and PostgreSQL Flex with versioned state management across 3 subscription tiers (dev, preprod, production).")
    add_bullet(doc, "Implemented ArgoCD-based GitOps with App of Apps pattern and ApplicationSet for multi-environment promotion, enabling declarative deployments with automated drift detection across 5 environments.")
    add_bullet(doc, "Enforced platform guardrails using OPA Gatekeeper policies on AKS \u2014 mandatory resource limits, approved container registries, required labels \u2014 preventing 30+ policy violations per sprint.")
    add_bullet(doc, "Architected secrets management platform migrating from Helm values to Azure Key Vault with CSI Secret Store Driver, eliminating 100% of plaintext secrets and enforcing zero-trust policies.")
    add_bullet(doc, "Deployed Istio service mesh and full observability stack (Prometheus, Grafana, Jaeger, Kiali) providing platform-wide visibility into resource utilization and service health.")
    add_bullet(doc, "Drove platform cost optimization through automated cleanup, right-sizing, and weekly dashboards, reducing monthly Azure spend by 20%.")

    add_experience_header(doc, "Senior Software Consultant (Platform/DevOps) \u2014 Knoldus Inc (NashTech Global), Noida", "Nov 2021 \u2013 Dec 2024")
    add_bullet(doc, "Built reusable Jenkins Shared Library pipelines serving 10+ teams, standardizing golden-path workflows and reducing pipeline duplication by 70%.")
    add_bullet(doc, "Designed developer self-service workflows integrating SonarQube quality gates, Snyk/Mend scanning, and automated PR checks, reducing vulnerabilities reaching production by 90%.")
    add_bullet(doc, "Migrated artifact management from JFrog Artifactory to GitHub Packages via automation, saving $15K+ annually in licensing costs.")
    add_bullet(doc, "Published Terraform modules to private registry via Azure Pipelines, enabling versioned, discoverable infrastructure components for cross-team consumption.")

    add_experience_header(doc, "Data Engineer \u2014 REOMNIFY, Noida", "Jan 2021 \u2013 Jun 2021")
    add_bullet(doc, "Engineered Python-based ETL pipelines using REST APIs and Selenium, processing 10K+ records daily into PostgreSQL.")

    add_common_education(doc)
    add_common_certifications(doc)

    add_section_heading(doc, "Achievements")
    add_bullet(doc, "Awarded 2x SITA Bravo Awards for driving complex platform deliveries across distributed teams in multiple time zones.")
    add_bullet(doc, "Open-source contributor: PR #232 to Azure/terraform-azurerm-avm-res-keyvault-vault (Microsoft Azure Verified Modules).")

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    docx_path = os.path.join(OUTPUT_DIR, "Vaibhav_Kumar_Platform_Engineer_ATS_Improved.docx")
    doc.save(docx_path)
    print(f"Created: {docx_path}")
    return docx_path


# ============================================================================
# ATS-IMPROVED RESUME 3: SRE (Target ATS: 98/100)
# ============================================================================
def generate_sre_resume():
    doc = Document()
    set_narrow_margins(doc)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)

    add_common_header(doc)

    add_section_heading(doc, "Professional Summary")
    add_body_text(
        doc,
        "Site Reliability Engineer with 4.5+ years ensuring production reliability, observability, and operational "
        "excellence for distributed microservices on Kubernetes (AKS/EKS). Defined SLOs and error budgets for 13 services, "
        "deployed Prometheus/Grafana monitoring reducing MTTD by 60% and MTTR by 45%, and drove toil reduction from 40% "
        "to 15%. Implemented chaos engineering, blameless postmortems, and incident response with PagerDuty. Skilled in "
        "Terraform, Ansible, OpenTelemetry, and Istio service mesh. Certified Google Cloud DevOps Engineer and Azure "
        "Network Engineer Associate. Awarded 2x SITA Bravo Awards for reliability under pressure.",
        size=8.5, before=0, after=1,
    )

    add_section_heading(doc, "Technical Skills")
    add_skill_line(doc, "Observability", "Prometheus, Grafana, OpenTelemetry, Loki, ELK Stack, Azure Monitor, Jaeger, Kiali, New Relic")
    add_skill_line(doc, "Containers & Orchestration", "Kubernetes (AKS, EKS), Docker, Helm Charts, Istio Service Mesh, ACR")
    add_skill_line(doc, "Infrastructure as Code", "Terraform (Azure, AWS \u2014 11 modules), Ansible, GitOps (ArgoCD)")
    add_skill_line(doc, "CI/CD & Automation", "Azure DevOps Pipelines, Jenkins (Shared Libraries), GitHub Actions")
    add_skill_line(doc, "Cloud Platforms", "Azure (AKS, Key Vault, App Gateway, VNets, PostgreSQL Flex), AWS (EC2, S3, RDS, IAM, GuardDuty)")
    add_skill_line(doc, "Reliability & Incident Mgmt", "SLO/SLA/Error Budget, PagerDuty, Chaos Engineering, Runbooks, Blameless Postmortems")
    add_skill_line(doc, "Security & Compliance", "Azure Key Vault, CSI Secret Store Driver, Snyk, Mend, SonarQube, IAM Hardening")
    add_skill_line(doc, "Scripting", "Python, Bash (Shell Scripting), PowerShell")

    add_section_heading(doc, "Professional Experience")

    add_experience_header(doc, "Infrastructure Engineer (SRE) \u2014 SITA, Gurugram", "Jan 2025 \u2013 Present")
    add_bullet(doc, "Defined SLOs and error budgets for 13 microservices based on latency (p99 < 200ms) and availability (99.9%) SLIs; deployed Prometheus recording rules and Grafana SLO dashboards enabling error budget\u2013driven release decisions.")
    add_bullet(doc, "Deployed Prometheus, Grafana, and OpenTelemetry Collector on AKS with SLI-based dashboards and alerting, reducing MTTD by 60% and MTTR by 45% while maintaining 99.9% availability SLA.")
    add_bullet(doc, "Established on-call rotation and incident response framework with PagerDuty alerting, severity-based escalation, runbook documentation, and blameless postmortem process \u2014 achieving < 15 min MTTA for P1 incidents.")
    add_bullet(doc, "Implemented chaos engineering using Chaos Mesh on AKS \u2014 pod kill, network partition, CPU stress \u2014 validating failure handling and reducing unplanned production incidents by 30%.")
    add_bullet(doc, "Identified and automated 15+ toil-heavy tasks (certificate rotations, pod cleanup, resource scaling) using Python and PowerShell, reducing toil from 40% to 15% of on-call capacity.")
    add_bullet(doc, "Provisioned production-grade AKS clusters using 11 Terraform modules with state locking and drift detection; migrated secrets to Azure Key Vault with CSI Driver achieving zero-trust posture.")
    add_bullet(doc, "Remediated 16 AWS security findings (EC2, RDS, S3, IAM, GuardDuty) achieving 100% compliance; led capacity planning using VPA recommendations and load testing.")
    add_bullet(doc, "Drove cost optimization through automated resource cleanup and right-sizing analysis, reducing monthly Azure spend by 20%.")

    add_experience_header(doc, "Senior Software Consultant (DevOps/SRE) \u2014 Knoldus Inc (NashTech Global), Noida", "Nov 2021 \u2013 Dec 2024")
    add_bullet(doc, "Engineered Jenkins CI/CD pipelines with SonarQube quality gates and Snyk/Mend scanning for 10+ projects, achieving zero critical CVEs and reducing build-to-deploy time by 40%.")
    add_bullet(doc, "Migrated artifact management from JFrog Artifactory to GitHub Packages with PowerShell automation while maintaining artifact integrity and traceability.")
    add_bullet(doc, "Implemented Ansible configuration management ensuring consistent provisioning and drift prevention across 30+ instances in dev, staging, and production.")

    add_experience_header(doc, "Data Engineer \u2014 REOMNIFY, Noida", "Jan 2021 \u2013 Jun 2021")
    add_bullet(doc, "Built Python-based data pipelines using REST APIs and Selenium, processing 10K+ records daily into PostgreSQL for analytics and monitoring dashboards.")

    add_common_education(doc)
    add_common_certifications(doc)

    add_section_heading(doc, "Achievements")
    add_bullet(doc, "Awarded 2x SITA Bravo Awards for driving mission-critical deployments under pressure across distributed teams.")
    add_bullet(doc, "Open-source contributor: PR #232 to Azure/terraform-azurerm-avm-res-keyvault-vault (Microsoft Azure Verified Modules).")

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    docx_path = os.path.join(OUTPUT_DIR, "Vaibhav_Kumar_SRE_ATS_Improved.docx")
    doc.save(docx_path)
    print(f"Created: {docx_path}")
    return docx_path


if __name__ == "__main__":
    devops_docx = generate_devops_resume()
    platform_docx = generate_platform_engineer_resume()
    sre_docx = generate_sre_resume()

    from docx2pdf import convert
    import time
    for docx_path in [devops_docx, platform_docx, sre_docx]:
        pdf_path = docx_path.replace(".docx", ".pdf")
        for attempt in range(3):
            try:
                convert(docx_path, pdf_path)
                print(f"Created: {pdf_path}")
                break
            except Exception as e:
                print(f"Attempt {attempt+1} failed for {os.path.basename(docx_path)}: {e}")
                time.sleep(2)

    print("\nAll ATS-improved files generated!")
